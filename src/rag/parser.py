"""
简历解析器 v2.0 — 纯规则提取，零LLM介入

P0 原则：
- 解析层（文本提取、字段识别）全部使用规则+正则，不使用LLM
- LLM 仅用于可选的字段标准化映射（normalizer.py）
- 每个提取字段标注来源（原文档行号/段落/分区）
- 保留原始文本副本，支持用户对比验证
"""
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from src.config import settings


# ============ 数据模型 ============

@dataclass
class ExtractionSource:
    """提取来源追踪 — 每个字段都知道自己来自哪里"""
    source_text: str = ""          # 原始文本片段
    source_line_start: int = 0     # 在 raw_text 中的起始行号（1-based）
    source_line_end: int = 0       # 在 raw_text 中的结束行号
    source_section: str = ""       # 所属分区（header/skills/projects/…）
    extraction_method: str = ""    # 提取方法（regex/keyword/pattern）
    confidence: float = 0.0        # 提取置信度 0.0-1.0


@dataclass
class ExtractionStats:
    """提取统计 — 用于用户验证"""
    total_skills_found: int = 0
    total_projects_found: int = 0
    total_achievements_found: int = 0
    total_education_found: int = 0
    total_work_found: int = 0
    high_confidence_count: int = 0
    medium_confidence_count: int = 0
    low_confidence_count: int = 0
    sections_identified: List[str] = field(default_factory=list)
    raw_text_length: int = 0
    raw_text_hash: str = ""


@dataclass
class ParsedResume:
    """结构化简历数据 — 每个字段带来源追踪"""
    # 基本信息
    name: str = ""
    name_source: Optional[ExtractionSource] = None
    email: str = ""
    email_source: Optional[ExtractionSource] = None
    phone: str = ""
    phone_source: Optional[ExtractionSource] = None

    # 结构化数据
    education: List[Dict[str, Any]] = field(default_factory=list)
    skills: List[Dict[str, Any]] = field(default_factory=list)
    projects: List[Dict[str, Any]] = field(default_factory=list)
    work_experience: List[Dict[str, Any]] = field(default_factory=list)
    achievements: List[Dict[str, Any]] = field(default_factory=list)

    # 原始文本（用于验证）
    raw_text: str = ""
    raw_text_hash: str = ""
    sections: Dict[str, str] = field(default_factory=dict)

    # 提取元数据
    extraction_stats: Optional[ExtractionStats] = None

    def get_verification_summary(self) -> str:
        """生成提取验证摘要 — 用户可对比原始文档"""
        stats = self.extraction_stats or ExtractionStats()
        lines = [
            f"=== 简历提取验证报告 ===",
            f"原始文档: {stats.raw_text_length} 字符 (SHA256: {stats.raw_text_hash[:12]}...)",
            f"识别分区: {', '.join(stats.sections_identified) if stats.sections_identified else '无'}",
            f"",
            f"提取结果:",
            f"  姓名: '{self.name}' (置信度: {self.name_source.confidence if self.name_source else 'N/A'})",
            f"  邮箱: '{self.email}'",
            f"  手机: '{self.phone}'",
            f"  技能: {stats.total_skills_found} 项",
            f"  项目: {stats.total_projects_found} 项",
            f"  成果: {stats.total_achievements_found} 项",
            f"  教育: {stats.total_education_found} 项",
            f"  工作: {stats.total_work_found} 项",
            f"",
            f"置信度分布:",
            f"  高置信度(≥0.8): {stats.high_confidence_count} 项",
            f"  中置信度(0.5-0.8): {stats.medium_confidence_count} 项",
            f"  低置信度(<0.5): {stats.low_confidence_count} 项",
        ]
        if stats.low_confidence_count > 0:
            lines.append(f"  ⚠️ 低置信度字段建议人工复核")
        return "\n".join(lines)

    def get_field_sources(self) -> List[Dict[str, Any]]:
        """获取所有字段的来源信息（用于前端展示）"""
        sources = []
        if self.name_source:
            sources.append({
                "field": "姓名", "value": self.name,
                "source_text": self.name_source.source_text[:100],
                "source_line": self.name_source.source_line_start,
                "section": self.name_source.source_section,
                "confidence": self.name_source.confidence,
            })
        for i, skill in enumerate(self.skills[:10]):
            src = skill.get("_source")
            if src:
                sources.append({
                    "field": f"技能#{i+1}", "value": skill.get("name", ""),
                    "source_text": src.source_text[:100] if hasattr(src, 'source_text') else "",
                    "source_line": src.source_line_start if hasattr(src, 'source_line_start') else 0,
                    "section": src.source_section if hasattr(src, 'source_section') else "",
                    "confidence": src.confidence if hasattr(src, 'confidence') else 0,
                })
        for i, proj in enumerate(self.projects[:5]):
            src = proj.get("_source")
            if src:
                sources.append({
                    "field": f"项目#{i+1}", "value": proj.get("name", ""),
                    "source_text": src.source_text[:100] if hasattr(src, 'source_text') else "",
                    "source_line": src.source_line_start if hasattr(src, 'source_line_start') else 0,
                    "section": src.source_section if hasattr(src, 'source_section') else "",
                    "confidence": src.confidence if hasattr(src, 'confidence') else 0,
                })
        return sources


@dataclass
class Document:
    """文档块"""
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunk_id: str = ""


# ============ 技能分类体系 ============

SKILL_CATEGORIES: Dict[str, List[str]] = {
    "programming": [
        "python", "java", "go", "golang", "rust", "c++", "cpp", "c#", "csharp",
        "typescript", "javascript", "js", "ts", "kotlin", "swift", "scala",
        "php", "ruby", "perl", "r", "matlab", "dart", "lua",
    ],
    "framework": [
        "react", "vue", "vue.js", "angular", "next.js", "nextjs", "nuxt",
        "fastapi", "django", "flask", "spring", "springboot", "express",
        "nestjs", "gin", "echo", "actix", "rocket", "laravel", "rails",
        "pytorch", "tensorflow", "keras", "langgraph", "langchain",
    ],
    "database": [
        "mysql", "postgresql", "postgres", "mongodb", "redis", "elasticsearch",
        "chromadb", "faiss", "milvus", "pinecone", "weaviate", "qdrant",
        "sqlite", "oracle", "sqlserver", "cassandra", "neo4j", "dynamodb",
    ],
    "devops": [
        "docker", "kubernetes", "k8s", "git", "github", "gitlab", "ci/cd",
        "jenkins", "github actions", "linux", "bash", "shell", "nginx",
        "apache", "terraform", "ansible", "prometheus", "grafana",
    ],
    "ai": [
        "machine learning", "deep learning", "nlp", "自然语言处理",
        "llm", "大语言模型", "rag", "agent", "transformer", "bert",
        "gpt", "embedding", "fine-tuning", "微调", "prompt engineering",
        "计算机视觉", "cv", "推荐系统", "数据挖掘",
    ],
    "cloud": [
        "aws", "azure", "gcp", "阿里云", "腾讯云", "华为云",
        "serverless", "lambda", "ec2", "s3", "rds",
    ],
    "design": [
        "figma", "sketch", "photoshop", "illustrator", "ui/ux",
        "用户体验", "交互设计", "原型设计",
    ],
}


def _classify_skill(name: str) -> Tuple[str, float]:
    """基于预定义分类体系对技能进行分类，返回(类别, 置信度)"""
    name_lower = name.lower().strip()
    for category, keywords in SKILL_CATEGORIES.items():
        for kw in keywords:
            if kw == name_lower or kw in name_lower or name_lower in kw:
                return category, 0.9
    return "other", 0.3


# ============ 主解析器 ============

class ResumeParser:
    """
    简历解析器 v2.0 — 纯规则提取，零LLM介入

    原则：
    1. 文本提取：直接读取，不做任何改写
    2. 字段识别：规则+正则，100%可追溯
    3. 每个字段标注来源行号、原始文本片段、提取方法、置信度
    4. 保留原始文本，用户可逐项对比验证
    """

    def __init__(self):
        self.supported = settings.get_supported_formats()

    def parse(self, file_path: str) -> ParsedResume:
        """解析简历文件，返回带来源追踪的结构化数据"""
        path = Path(file_path)
        ext = path.suffix.lower().lstrip(".")
        if ext not in self.supported:
            raise ValueError(f"不支持的格式: .{ext}，支持: {self.supported}")
        if not path.exists():
            raise FileNotFoundError(f"简历文件不存在: {file_path}")

        # Step 1: 提取原始文本（零改写）
        raw_text = self._extract_text(file_path, ext)
        raw_hash = hashlib.sha256(raw_text.encode('utf-8')).hexdigest()

        # Step 2: 分行索引（用于来源追踪）
        raw_lines = raw_text.split('\n')

        # Step 3: 按段落分割
        paragraphs = self._split_paragraphs(raw_text)

        # Step 4: 识别简历分区
        sections = self._identify_sections(paragraphs)

        # Step 5: 构建结构化数据（每个字段带来源）
        parsed = ParsedResume(
            raw_text=raw_text,
            raw_text_hash=raw_hash,
            sections=sections,
        )

        # 基本信息（带来源追踪）
        name, name_src = self._extract_name_with_source(raw_text, raw_lines)
        parsed.name = name
        parsed.name_source = name_src

        email, email_src = self._extract_email_with_source(raw_text, raw_lines)
        parsed.email = email
        parsed.email_source = email_src

        phone, phone_src = self._extract_phone_with_source(raw_text, raw_lines)
        parsed.phone = phone
        parsed.phone_source = phone_src

        # 结构化模块（全部带来源追踪）
        parsed.skills = self._extract_skills_v2(sections, raw_lines)
        parsed.projects = self._extract_projects_v2(sections, raw_lines)
        parsed.work_experience = self._extract_work_experience_v2(sections, raw_lines)
        parsed.education = self._extract_education_v2(sections, raw_lines)
        parsed.achievements = self._extract_achievements_v2(sections, parsed.projects, raw_lines)

        # 生成提取统计
        stats = ExtractionStats(
            total_skills_found=len(parsed.skills),
            total_projects_found=len(parsed.projects),
            total_achievements_found=len(parsed.achievements),
            total_education_found=len(parsed.education),
            total_work_found=len(parsed.work_experience),
            sections_identified=list(sections.keys()),
            raw_text_length=len(raw_text),
            raw_text_hash=raw_hash,
        )

        # 统计置信度分布
        for items in [parsed.skills, parsed.projects, parsed.achievements,
                       parsed.education, parsed.work_experience]:
            for item in items:
                src = item.get("_source")
                if src and hasattr(src, 'confidence'):
                    if src.confidence >= 0.8:
                        stats.high_confidence_count += 1
                    elif src.confidence >= 0.5:
                        stats.medium_confidence_count += 1
                    else:
                        stats.low_confidence_count += 1

        # 姓名也计入
        if parsed.name_source and parsed.name_source.confidence >= 0.8:
            stats.high_confidence_count += 1
        elif parsed.name_source and parsed.name_source.confidence >= 0.5:
            stats.medium_confidence_count += 1

        parsed.extraction_stats = stats
        return parsed

    def to_documents(self, parsed: ParsedResume) -> List[Document]:
        """将结构化简历转为文档块列表（用于向量化），保留来源信息"""
        documents = []

        for i, skill in enumerate(parsed.skills):
            content = f"技能: {skill.get('name', '')}"
            if skill.get('category'):
                content += f" (类别: {skill['category']})"
            meta = {
                "type": "skills", "index": i,
                "name": skill.get("name", ""),
                "category": skill.get("category", ""),
            }
            src = skill.get("_source")
            if src:
                meta["source_text"] = src.source_text[:200] if hasattr(src, 'source_text') else ""
                meta["confidence"] = src.confidence if hasattr(src, 'confidence') else 0
            documents.append(Document(content=content, metadata=meta, chunk_id=f"skill_{i}"))

        for i, proj in enumerate(parsed.projects):
            parts = [f"项目: {proj.get('name', '')}"]
            if proj.get('role'):
                parts.append(f"角色: {proj['role']}")
            if proj.get('tech_stack'):
                techs = proj['tech_stack'] if isinstance(proj['tech_stack'], list) else [proj['tech_stack']]
                parts.append(f"技术栈: {', '.join(techs)}")
            if proj.get('key_result'):
                parts.append(f"关键成果: {proj['key_result']}")
            meta = {
                "type": "projects", "index": i,
                "name": proj.get("name", ""),
                "role": proj.get("role", ""),
            }
            src = proj.get("_source")
            if src:
                meta["source_text"] = src.source_text[:200] if hasattr(src, 'source_text') else ""
                meta["confidence"] = src.confidence if hasattr(src, 'confidence') else 0
            documents.append(Document(content="\n".join(parts), metadata=meta, chunk_id=f"project_{i}"))

        for i, work in enumerate(parsed.work_experience):
            parts = [f"工作经历: {work.get('company', '')} - {work.get('position', '')}"]
            if work.get('duration'):
                parts.append(f"时间: {work['duration']}")
            meta = {"type": "projects", "index": i, "company": work.get("company", "")}
            src = work.get("_source")
            if src:
                meta["confidence"] = src.confidence if hasattr(src, 'confidence') else 0
            documents.append(Document(content="\n".join(parts), metadata=meta, chunk_id=f"work_{i}"))

        for i, ach in enumerate(parsed.achievements):
            meta = {"type": "achievements", "index": i}
            src = ach.get("_source")
            if src:
                meta["source_text"] = src.source_text[:200] if hasattr(src, 'source_text') else ""
                meta["confidence"] = src.confidence if hasattr(src, 'confidence') else 0
            documents.append(Document(
                content=f"成果: {ach.get('description', '')}",
                metadata=meta, chunk_id=f"achievement_{i}"
            ))

        for i, edu in enumerate(parsed.education):
            parts = [f"教育: {edu.get('school', '')} - {edu.get('degree', '')}"]
            if edu.get('major'):
                parts.append(f"专业: {edu['major']}")
            if edu.get('time'):
                parts.append(f"时间: {edu['time']}")
            meta = {"type": "education", "index": i, "school": edu.get("school", "")}
            src = edu.get("_source")
            if src:
                meta["confidence"] = src.confidence if hasattr(src, 'confidence') else 0
            documents.append(Document(content="\n".join(parts), metadata=meta, chunk_id=f"education_{i}"))

        return documents

    # ===== 文本提取（零改写）=====

    def _extract_text(self, file_path: str, ext: str) -> str:
        """直接提取文本，不做任何改写"""
        if ext == "pdf":
            return self._extract_pdf(file_path)
        elif ext == "docx":
            return self._extract_docx(file_path)
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()

    def _extract_pdf(self, file_path: str) -> str:
        """PyMuPDF 直接提取"""
        try:
            import fitz
            doc = fitz.open(file_path)
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text("text"))
            doc.close()
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("需要安装 PyMuPDF: pip install PyMuPDF")

    def _extract_docx(self, file_path: str) -> str:
        """python-docx 提取，失败回退 XML"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)
            result = "\n\n".join(text_parts)
            if len(result) < 100:
                result = self._extract_docx_xml(file_path)
            return result
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")

    def _extract_docx_xml(self, file_path: str) -> str:
        """DOCX XML 直接解析"""
        import zipfile
        import xml.etree.ElementTree as ET
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        with zipfile.ZipFile(file_path) as z:
            with z.open('word/document.xml') as f:
                tree = ET.parse(f)
        root = tree.getroot()
        paragraphs = []
        for para_elem in root.findall('.//w:p', ns):
            texts = []
            for t in para_elem.findall('.//w:t', ns):
                if t.text:
                    texts.append(t.text)
            if texts:
                line = ''.join(texts).strip()
                if line:
                    paragraphs.append(line)
        return '\n\n'.join(paragraphs)

    def _split_paragraphs(self, text: str) -> List[str]:
        """按段落分割"""
        paras = re.split(r'\n\s*\n', text)
        return [p.strip() for p in paras if p.strip()]

    def _identify_sections(self, paragraphs: List[str]) -> Dict[str, str]:
        """识别简历分区（严格版 - 精确匹配分区标题行）"""
        sections = {}
        current_section = "header"
        current_content: List[str] = []

        # 严格分区标题: 只匹配作为单独一行的标题
        # 格式: "XX经历"、"XX背景"、"XX技能"、"XX项目" 等
        section_patterns = [
            ("education", re.compile(
                r'^(?:教育背景|教育经历|学历|education|academic|主修课程)'
                r'(?:[：:].*)?\s*$', re.IGNORECASE)),
            ("skills", re.compile(
                r'^(?:专业技能|技能|技术栈|技术能力|skills|technologies)'
                r'(?:[：:].*)?\s*$', re.IGNORECASE)),
            ("work", re.compile(
                r'^(?:工作经历|工作经验|实习经历|科研经历|科研项目|'
                r'work experience|experience|employment)'
                r'(?:[：:].*)?\s*$', re.IGNORECASE)),
            ("projects", re.compile(
                r'^(?:项目经历|项目经验|实习项目|项目|projects|project experience)'
                r'(?:[：:].*)?\s*$', re.IGNORECASE)),
            ("achievements", re.compile(
                r'^(?:获奖|荣誉证书|证书|achievements|awards|honors|'
                r'论文发表|论文|发表)'
                r'(?:[：:].*)?\s*$', re.IGNORECASE)),
        ]

        for para in paragraphs:
            first_line = para.strip().split("\n")[0].strip()
            first_line_clean = re.sub(r'^#+\s*', '', first_line)

            matched = False
            # 只匹配短行（分区标题通常 < 30 字符）
            if len(first_line_clean) <= 30:
                for section_name, pattern in section_patterns:
                    if pattern.match(first_line_clean):
                        if current_content:
                            sections[current_section] = "\n".join(current_content)
                        current_section = section_name
                        current_content = []
                        lines = para.strip().split("\n")
                        remaining = [l.strip() for l in lines[1:] if l.strip() and not l.strip().startswith("#")]
                        if remaining:
                            current_content.extend(remaining)
                        matched = True
                        break

            if not matched:
                current_content.append(para)

        if current_content:
            sections[current_section] = "\n".join(current_content)

        return sections

    # ===== 字段提取（带来源追踪）=====

    def _find_line_range(self, text_snippet: str, raw_lines: List[str]) -> Tuple[int, int]:
        """在 raw_lines 中查找文本片段的大致行号范围"""
        snippet_first = text_snippet.strip().split('\n')[0][:30] if text_snippet else ""
        for i, line in enumerate(raw_lines):
            if snippet_first and snippet_first in line:
                # 向后查找结束
                end = i
                for j in range(i, min(i + 10, len(raw_lines))):
                    remaining = text_snippet.strip().split('\n')[-1][-30:] if '\n' in text_snippet else text_snippet[-30:]
                    if remaining and remaining[:20] in raw_lines[j]:
                        end = j
                        break
                return i + 1, end + 1  # 1-based
        return 0, 0

    def _extract_name_with_source(self, text: str, raw_lines: List[str]) -> Tuple[str, Optional[ExtractionSource]]:
        """提取姓名 + 来源追踪"""
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        for idx, line in enumerate(lines[:10]):
            if line.startswith("#"):
                continue
            # "姓名" 标签
            name_match = re.search(r'(?:姓名|名字)[：:]\s*(\S{2,4})', line)
            if name_match:
                name = name_match.group(1)
                src = ExtractionSource(
                    source_text=line,
                    source_line_start=idx + 1,
                    source_line_end=idx + 1,
                    source_section="header",
                    extraction_method="label_match",
                    confidence=0.95,
                )
                return name, src
            # 中文名模式
            parts = re.split(r'\s*[|｜\t]{2,}\s*', line)
            for part in parts:
                part = part.strip()
                if not part or len(part) > 15:
                    continue
                if any(kw in part.lower() for kw in [
                    "简历", "resume", "cv", "电话", "邮箱", "email", "手机",
                    "地址", "github", "linkedin", "博客", "blog", "姓名",
                ]):
                    continue
                if "@" in part or re.search(r'\d{8,}', part):
                    continue
                chinese_name = re.search(r'([一-鿿]{2,4})', part)
                if chinese_name:
                    name = chinese_name.group(1)
                    src = ExtractionSource(
                        source_text=part,
                        source_line_start=idx + 1,
                        source_line_end=idx + 1,
                        source_section="header",
                        extraction_method="chinese_name_pattern",
                        confidence=0.85,
                    )
                    return name, src

        return "未知", ExtractionSource(
            source_text="", source_section="header",
            extraction_method="none", confidence=0.0,
        )

    def _extract_email_with_source(self, text: str, raw_lines: List[str]) -> Tuple[str, Optional[ExtractionSource]]:
        """提取邮箱 + 来源"""
        match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        if match:
            email = match.group(0)
            line_range = self._find_line_range(email, raw_lines)
            src = ExtractionSource(
                source_text=email,
                source_line_start=line_range[0],
                source_line_end=line_range[1],
                source_section="header",
                extraction_method="email_regex",
                confidence=0.98,
            )
            return email, src
        return "", ExtractionSource(source_text="", extraction_method="none", confidence=0.0)

    def _extract_phone_with_source(self, text: str, raw_lines: List[str]) -> Tuple[str, Optional[ExtractionSource]]:
        """提取手机号 + 来源"""
        match = re.search(r'1[3-9]\d{9}', text)
        if match:
            phone = match.group(0)
            line_range = self._find_line_range(phone, raw_lines)
            src = ExtractionSource(
                source_text=phone,
                source_line_start=line_range[0],
                source_line_end=line_range[1],
                source_section="header",
                extraction_method="phone_regex",
                confidence=0.98,
            )
            return phone, src
        return "", ExtractionSource(source_text="", extraction_method="none", confidence=0.0)

    def _extract_skills_v2(self, sections: Dict[str, str], raw_lines: List[str]) -> List[Dict[str, Any]]:
        """
        增强版技能提取：
        - 更好的分词（处理多词技能如 "Machine Learning"）
        - 基于预定义分类体系的分类
        - 每个技能带来源追踪和置信度
        """
        skills = []
        skill_text = sections.get("skills", "")
        if not skill_text:
            return skills

        # 按行处理
        skill_lines = skill_text.split('\n')
        for line in skill_lines:
            line = line.strip()
            if not line:
                continue

            # 尝试多种分隔符
            # 按逗号/中文逗号/顿号/竖线分割
            tokens = re.split(r'[,，、|/·•·]\s*', line)
            if len(tokens) == 1:
                # 没有分隔符，整行作为一个技能
                tokens = [line]

            for token in tokens:
                token = token.strip()
                # 清理：去掉列表符号和多余的空格
                token = re.sub(r'^[\-\*\•\·]\s*', '', token)
                token = re.sub(r'^[0-9]+[\.\)、]\s*', '', token)
                token = token.strip()

                if not token or len(token) > 30 or len(token) < 1:
                    continue

                # 过滤明显不是技能的文本
                if any(kw in token.lower() for kw in ["熟练", "掌握", "了解", "精通", "熟悉"]):
                    # 提取熟练度描述中的技能名
                    inner = re.sub(r'(?:熟练|掌握|了解|精通|熟悉)\s*[：:度程度]?\s*', '', token)
                    if inner.strip():
                        token = inner.strip()

                # 分类 + 置信度
                category, confidence = _classify_skill(token)

                # 来源追踪
                line_range = self._find_line_range(token, raw_lines)

                skills.append({
                    "name": token,
                    "category": category,
                    "level": "",
                    "_source": ExtractionSource(
                        source_text=token,
                        source_line_start=line_range[0],
                        source_line_end=line_range[1],
                        source_section="skills",
                        extraction_method="delimiter_split",
                        confidence=confidence,
                    ),
                })

        return skills

    def _extract_projects_v2(self, sections: Dict[str, str], raw_lines: List[str]) -> List[Dict[str, Any]]:
        """增强版项目提取：多策略分割，适配多种简历格式"""
        projects = []
        project_text = sections.get("projects", "")
        if not project_text:
            return projects

        # 项目分割：合并两种边界格式，一次识别所有项目，避免顺序互斥导致漏分
        #   边界A: "项目名 + 4+空白 + 日期"（如 "ResuMatch AI ... 2026.05 - 至今"）
        #   边界B: "项目名 | 角色 [公司]"（如 "视觉康复随访管理系统 | 全栈开发工程师  XX科技"）
        # 用 | 交替合成为单一正则，任意匹配即切分，保证两种格式的项目都能被识别
        project_blocks = re.split(
            r'\n(?=[A-Za-z一-鿿（(](?:[^\n]{3,120}\s{4,}\d{4}[.\-]'
            r'|[^\n]{5,100}\s*\|\s*[^\n]{1,30}))',
            project_text
        )

        for block in project_blocks:
            block = block.strip()
            if not block or len(block) < 15:
                continue

            lines = block.split("\n")
            first_line = lines[0].strip() if lines else "未知项目"
            # 清理项目名中的编号前缀
            name = re.sub(r'^[\d]+[\.\)、]\s*', '', first_line)
            # 清理尾部长空白及日期
            name = re.sub(r'\s{4,}.*$', '', name)
            name = name.strip()
            # 从 "项目名 | 角色   公司" 首行中拆分出角色（| 分隔）
            role = ""
            if "|" in name:
                parts = name.split("|", 1)
                name = parts[0].strip()
                role_raw = parts[1].strip()
                # 角色通常为 2-8 字，截取 "|" 后的第一个角色段（排除尾随的公司名）
                role_match = re.match(r'^([^\s]{2,12}?)(?:\s{2,}.*)?$', role_raw)
                if role_match:
                    role = role_match.group(1).strip()
            if len(name) > 120:
                name = name[:120]
            # 跳过明显不是项目名的块
            if re.match(r'^(?:项目简介|设计|主要工作|简历|技能|工具|语言|基于前后|负责|封装|搭建|引入|实现)', name):
                continue
            if len(name) < 5:
                continue

            # 提取角色（若首行未含 |，则从 "角色/岗位/职位: xxx" 提取）
            if not role:
                role_match = re.search(r'(?:角色|岗位|职位|role)[：:]\s*([^\n]{2,15})', block, re.IGNORECASE)
                if role_match:
                    role = role_match.group(1).strip()

            # 提取时间段
            time_period = ""
            time_match = re.search(
                r'(?:时间|日期|period|duration)[：:]\s*([^\n]{3,30})',
                block, re.IGNORECASE
            )
            if time_match:
                time_period = time_match.group(1).strip()
            else:
                # 兼容 "2026.05 – 至今" / "2026.05 - 2026.09" / "2024.03-2024.09" 等格式
                date_match = re.search(
                    r'(\d{4}[./年]\d{1,2}[月]?\s*[-–—至到]\s*'
                    r'(?:\d{4}[./年]\d{1,2}[月]?|\d{4}|\s*至今|现在|present))',
                    block, re.IGNORECASE
                )
                if date_match:
                    time_period = date_match.group(1)

            # 提取技术栈: 只认显式标签（"技术栈"/"tech stack"），
            # 裸 "技术：" 仅在非中文词尾时视为标签，避免误抓 "6项RAG增强技术：" 这类正文
            tech_stack = []
            tech_match = re.findall(
                r'(?:技术栈|tech stack|technologies|(?<![一-鿿])技术)[：:]\s*(.+?)(?:\n|$)',
                block, re.IGNORECASE
            )
            if tech_match:
                tech_stack = [t.strip() for t in re.split(r'[,，、/]', tech_match[0]) if t.strip()]
            # 如果没找到\"技术栈\"标签，从内容中提取技术关键词
            if not tech_stack:
                tech_keywords = re.findall(
                    r'(LangGraph|FastAPI|ChromaDB|DeepSeek|Transformer|BERT|BART|LoRA|'
                    r'HyDE|Self-Query|Cross-Encoder|bge|RAG|LLaVA|CoreNLP|SAM|FG-CLIP|'
                    r'Python|C\+\+|Docker|Git|Linux|MySQL|PostgreSQL|Redis|'
                    r'FAISS|Milvus|Streamlit|React|Vue|TypeScript)',
                    block, re.IGNORECASE
                )
                tech_stack = list(dict.fromkeys(tech_keywords))  # 去重保序

            # 提取角色（若前面已从 "| 角色" 提取到则保留）
            if not role:
                role_match = re.search(r'(?:角色|岗位|职位|role)[：:]\s*(.+?)(?:\n|$)', block, re.IGNORECASE)
                if role_match:
                    role = role_match.group(1).strip()

            # 提取量化成果
            key_result = self._extract_key_result_v2(block)
            # 如果没提取到，用\"项目简介\"行+量化行作为替代
            if not key_result:
                intro_match = re.search(r'项目简介[：:]\s*(.+?)(?:\n|$)', block)
                if intro_match:
                    key_result = intro_match.group(1).strip()[:300]

            # 来源追踪
            line_range = self._find_line_range(name, raw_lines)
            has_data = bool(tech_stack or role or key_result)
            confidence = 0.85 if has_data else 0.5

            projects.append({
                "name": name,
                "description": block[:800],
                "role": role,
                "tech_stack": tech_stack,
                "time_period": time_period,
                "key_result": key_result,
                "_source": ExtractionSource(
                    source_text=name,
                    source_line_start=line_range[0],
                    source_line_end=line_range[1],
                    source_section="projects",
                    extraction_method="section_header_split",
                    confidence=confidence,
                ),
            })

        # 合并小块：只有当后续块"不像独立项目"（无日期/角色/管道，短名且无项目关键词）才视为延续
        # 避免误吞自带项目头信号的块（如 "西湖大学张紫阳实验室 ... 2024.03-2024.09" 是独立项目）
        merged = []
        for proj in projects:
            has_header_signal = bool(proj.get("time_period") or proj.get("role") or "|" in proj.get("name", ""))
            if (merged and not has_header_signal
                    and len(proj["name"]) < 20 and not re.search(r'[项目系统平台助手工具]', proj["name"])):
                # 这是上一块的延续，追加描述
                merged[-1]["description"] += "\n" + proj["name"] + "\n" + proj["description"]
            else:
                # 去掉\"项目简介\"等非项目名块（它们属于上一个项目)
                if not re.match(r'^(?:项目简介|设计|简历|6项|5项)', proj["name"]):
                    merged.append(proj)
                elif merged:
                    merged[-1]["description"] += "\n" + proj["name"] + "\n" + proj["description"]

        return merged

    def _extract_key_result_v2(self, text: str) -> str:
        """增强版量化成果提取"""
        patterns = [
            # 百分比相关
            r'[^。\n]{0,80}(?:\d+%|提升\d+|降低\d+|优化\d+|减少\d+|增长\d+)[^。\n]{0,80}',
            # 数值区间
            r'[^。\n]{0,80}(?:从\d+到\d+|提高\d+倍|节省\d+|缩短\d+)[^。\n]{0,80}',
            # 具体指标
            r'[^。\n]{0,80}(?:QPS|延迟|响应时间|准确率|召回率|转化率|覆盖率)[^。\n]{0,80}',
        ]
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                return matches[0].strip()
        return ""

    def _extract_work_experience_v2(self, sections: Dict[str, str], raw_lines: List[str]) -> List[Dict[str, Any]]:
        """工作经历提取 + 来源追踪"""
        work = []
        work_text = sections.get("work", "")
        if not work_text:
            return work

        blocks = re.split(r'\n(?=[A-Za-z一-鿿][^\n]{0,50}(?:公司|科技|集团|有限|实验室|研究院))', work_text)
        for block in blocks:
            block = block.strip()
            if not block or len(block) < 5:
                continue
            lines = block.split("\n")
            company = lines[0].strip() if lines else "未知公司"

            position = ""
            duration = ""
            pos_match = re.search(r'(?:职位|岗位|角色)[：:]\s*(.+?)(?:\n|$)', block)
            if pos_match:
                position = pos_match.group(1).strip()
            time_match = re.search(r'(?:时间|日期|period)[：:]\s*(.+?)(?:\n|$)', block)
            if time_match:
                duration = time_match.group(1).strip()

            line_range = self._find_line_range(company, raw_lines)

            work.append({
                "company": company,
                "position": position,
                "duration": duration,
                "description": block[:300],
                "_source": ExtractionSource(
                    source_text=company,
                    source_line_start=line_range[0],
                    source_line_end=line_range[1],
                    source_section="work",
                    extraction_method="section_split",
                    confidence=0.7 if position else 0.5,
                ),
            })

        return work

    def _extract_education_v2(self, sections: Dict[str, str], raw_lines: List[str]) -> List[Dict[str, Any]]:
        """教育背景提取 + 来源追踪"""
        education = []
        edu_text = sections.get("education", "")
        if not edu_text:
            return education

        lines = edu_text.split("\n")
        for line in lines:
            line = line.strip()
            if not line or len(line) < 5:
                continue

            # 尝试提取学校、学位、专业、时间
            school = line
            degree = ""
            major = ""
            time_period = ""

            # 学位匹配
            degree_match = re.search(r'(?:学士|硕士|博士|本科|研究生|MBA|PhD|Bachelor|Master)', line)
            if degree_match:
                degree = degree_match.group(0)
                school = line.replace(degree, "").strip()

            # 专业匹配
            major_match = re.search(
                r'(?:计算机|软件|数据|人工智能|电子|通信|机械|金融|管理|数学|物理|化学|生物|'
                r'Computer Science|Software Engineering|Data Science|AI)',
                line, re.IGNORECASE
            )
            if major_match:
                major = major_match.group(0)

            # 时间匹配
            time_match = re.search(r'(\d{4}[./年]\d{1,2}[月]?\s*[-–—至到]\s*\d{4}[./年]\d{1,2}[月]?)', line)
            if time_match:
                time_period = time_match.group(1)

            line_range = self._find_line_range(line, raw_lines)

            education.append({
                "school": school,
                "degree": degree,
                "major": major,
                "time": time_period,
                "_source": ExtractionSource(
                    source_text=line,
                    source_line_start=line_range[0],
                    source_line_end=line_range[1],
                    source_section="education",
                    extraction_method="line_parse",
                    confidence=0.8 if (degree or major) else 0.5,
                ),
            })

        return education

    def _extract_achievements_v2(
        self, sections: Dict[str, str], projects: List[Dict[str, Any]], raw_lines: List[str]
    ) -> List[Dict[str, Any]]:
        """成就提取 + 来源追踪"""
        achievements = []

        # 从成就分区提取
        ach_text = sections.get("achievements", "")
        if ach_text:
            for line in ach_text.split("\n"):
                line = line.strip()
                if line and len(line) > 5:
                    line = re.sub(r'^[\-\*\•\·]\s*', '', line)
                    line = re.sub(r'^[0-9]+[\.\)、]\s*', '', line)
                    line_range = self._find_line_range(line, raw_lines)
                    achievements.append({
                        "description": line,
                        "_source": ExtractionSource(
                            source_text=line,
                            source_line_start=line_range[0],
                            source_line_end=line_range[1],
                            source_section="achievements",
                            extraction_method="line_parse",
                            confidence=0.8,
                        ),
                    })

        # 从项目中提取量化成果
        for proj in projects:
            if proj.get("key_result"):
                achievements.append({
                    "description": f"[{proj.get('name', '项目')}] {proj['key_result']}",
                    "_source": ExtractionSource(
                        source_text=proj["key_result"],
                        source_section="projects",
                        extraction_method="key_result_extraction",
                        confidence=0.75,
                    ),
                })

        return achievements

    # ===== 向后兼容包装器（旧API → 新API）=====

    def _extract_name(self, text: str) -> str:
        """[兼容] 旧版姓名提取"""
        name, _ = self._extract_name_with_source(text, text.split('\n'))
        return name

    def _extract_email(self, text: str) -> str:
        """[兼容] 旧版邮箱提取"""
        email, _ = self._extract_email_with_source(text, text.split('\n'))
        return email

    def _extract_phone(self, text: str) -> str:
        """[兼容] 旧版手机号提取"""
        phone, _ = self._extract_phone_with_source(text, text.split('\n'))
        return phone

    def _extract_skills(self, sections: Dict[str, str]) -> List[Dict[str, Any]]:
        """[兼容] 旧版技能提取 — 内部调用 v2 并去除 _source"""
        raw_lines = sections.get("skills", "").split('\n')
        results = self._extract_skills_v2(sections, raw_lines)
        return [{k: v for k, v in r.items() if k != "_source"} for r in results]

    def _extract_projects(self, sections: Dict[str, str]) -> List[Dict[str, Any]]:
        """[兼容] 旧版项目提取"""
        raw_lines = sections.get("projects", "").split('\n')
        results = self._extract_projects_v2(sections, raw_lines)
        return [{k: v for k, v in r.items() if k != "_source"} for r in results]

    def _extract_work_experience(self, sections: Dict[str, str]) -> List[Dict[str, Any]]:
        """[兼容] 旧版工作经历提取"""
        raw_lines = sections.get("work", "").split('\n')
        results = self._extract_work_experience_v2(sections, raw_lines)
        return [{k: v for k, v in r.items() if k != "_source"} for r in results]

    def _extract_education(self, sections: Dict[str, str]) -> List[Dict[str, Any]]:
        """[兼容] 旧版教育背景提取"""
        raw_lines = sections.get("education", "").split('\n')
        results = self._extract_education_v2(sections, raw_lines)
        return [{k: v for k, v in r.items() if k != "_source"} for r in results]

    def _extract_achievements(
        self, sections: Dict[str, str], projects: List[Dict[str, Any]]
    ) -> List[Dict[str, str]]:
        """[兼容] 旧版成就提取"""
        raw_lines = sections.get("achievements", "").split('\n')
        # 需要给 projects 加上 _source 兼容
        compat_projects = []
        for p in projects:
            if "_source" not in p:
                p = {**p, "_source": None}
            compat_projects.append(p)
        results = self._extract_achievements_v2(sections, compat_projects, raw_lines)
        return [{k: v for k, v in r.items() if k != "_source"} for r in results]
